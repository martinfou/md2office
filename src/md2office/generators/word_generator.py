"""
Word Document Generator

Implements Epic 2: Word Document Conversion
Generates Microsoft Word (.docx) documents from AST.
"""

from typing import Dict, Any, Optional, List, Union
from io import BytesIO
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Create a dummy RGBColor class for type hints when docx is not available
    class RGBColor:  # type: ignore
        def __init__(self, r: int, g: int, b: int):
            pass

from ..parser.ast_builder import ASTNode, NodeType
from ..router.content_router import FormatGenerator, OutputFormat
from ..errors import ConversionError, FileError
from ..styling.style import StylePreset, get_style_preset
from .inline_formatter import InlineFormatter


class WordGenerator(FormatGenerator):
    """
    Generates Word (.docx) documents from AST.
    
    Implements Story 2.1: Word Document Generator Core
    """
    
    def __init__(self):
        """Initialize Word generator."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word generation. "
                "Install with: pip install python-docx"
            )
        self.document: Optional[Document] = None
        self.current_style_preset: Optional[StylePreset] = None
        self.inline_formatter = InlineFormatter()
    
    def generate(self, ast: ASTNode, options: Dict[str, Any]) -> bytes:
        """
        Generate Word document from AST.
        
        Args:
            ast: Root AST node
            options: Generation options
            
        Returns:
            Generated Word document as bytes
            
        Raises:
            ConversionError: If generation fails
        """
        try:
            # Initialize document
            self.document = Document()
            
            # Get style preset
            style_name = options.get('style', 'default')
            self.current_style_preset = get_style_preset(style_name)
            
            # Set document metadata
            self._set_document_metadata(ast, options)
            
            # Process AST nodes
            self._process_node(ast, options)
            
            # Add table of contents if requested
            if options.get('table_of_contents', False):
                self._add_table_of_contents()
            
            # Save to bytes
            output = BytesIO()
            self.document.save(output)
            return output.getvalue()
        
        except Exception as e:
            raise ConversionError(
                f"Failed to generate Word document: {str(e)}",
                format="word",
                stage="generation"
            ) from e
    
    def get_file_extension(self) -> str:
        """Get file extension for Word format."""
        return ".docx"
    
    def _set_document_metadata(self, ast: ASTNode, options: Dict[str, Any]):
        """Set document metadata (title, author, etc.)."""
        core_props = self.document.core_properties
        
        # Get title from AST metadata or first H1
        title = ast.metadata.get('title')
        if not title:
            first_h1 = self._find_first_heading(ast, 1)
            if first_h1:
                title = first_h1.content
        
        if title:
            core_props.title = title
        
        # Get author from metadata or options
        author = ast.metadata.get('author') or options.get('author')
        if author:
            core_props.author = author
        
        # Get subject from metadata
        subject = ast.metadata.get('subject')
        if subject:
            core_props.subject = subject
        
        # Get keywords from metadata
        keywords = ast.metadata.get('keywords')
        if keywords:
            core_props.keywords = keywords
    
    def _process_node(self, node: ASTNode, options: Dict[str, Any]):
        """
        Process AST node and add to document.
        
        Args:
            node: AST node to process
            options: Generation options
        """
        if node.node_type == NodeType.DOCUMENT:
            # Process all children of document
            for child in node.children:
                self._process_node(child, options)
        
        elif node.node_type == NodeType.SECTION:
            # Process section (may add page break)
            if options.get('page_breaks', False) and node.level == 1:
                # Add page break before H1 sections
                if self.document.paragraphs:
                    self.document.paragraphs[-1].runs[-1].add_break(6)  # Page break
            
            # Process section children
            for child in node.children:
                self._process_node(child, options)
        
        elif node.node_type == NodeType.HEADING:
            self._add_heading(node, options)
        
        elif node.node_type == NodeType.PARAGRAPH:
            self._add_paragraph(node, options)
        
        elif node.node_type == NodeType.LIST:
            self._add_list(node, options)
        
        elif node.node_type == NodeType.TABLE:
            self._add_table(node, options)
        
        elif node.node_type == NodeType.CODE_BLOCK:
            self._add_code_block(node, options)
        
        elif node.node_type == NodeType.BLOCKQUOTE:
            self._add_blockquote(node, options)
        
        elif node.node_type == NodeType.HORIZONTAL_RULE:
            self._add_horizontal_rule()
        
        elif node.node_type == NodeType.IMAGE:
            self._add_image(node, options)
        
        # Process children recursively
        for child in node.children:
            if child.node_type not in [NodeType.HEADING, NodeType.PARAGRAPH, 
                                      NodeType.LIST, NodeType.TABLE, 
                                      NodeType.CODE_BLOCK, NodeType.BLOCKQUOTE,
                                      NodeType.HORIZONTAL_RULE, NodeType.IMAGE]:
                self._process_node(child, options)
    
    def _add_heading(self, node: ASTNode, options: Dict[str, Any]):
        """Add heading to document."""
        level = node.level or 1
        content = node.content
        
        # Use Word's built-in heading styles
        style_name = f'Heading {level}'
        
        # Create paragraph with heading style
        paragraph = self.document.add_heading(content, level=level)
        
        # Apply custom styling if preset is available
        if self.current_style_preset:
            heading_style = self.current_style_preset.get_heading_style(level)
            
            # Apply font size
            for run in paragraph.runs:
                run.font.size = Pt(heading_style.font.size)
                run.font.bold = heading_style.font.weight == "bold"
                if heading_style.font.color:
                    run.font.color.rgb = self._parse_color(heading_style.font.color)
        
        # Add bookmark for navigation
        if options.get('bookmarks', True):
            self._add_bookmark(paragraph, content)
    
    def _add_paragraph(self, node: ASTNode, options: Dict[str, Any]):
        """Add paragraph to document."""
        paragraph = self.document.add_paragraph()
        
        # Apply paragraph style
        if self.current_style_preset:
            para_style = self.current_style_preset.paragraph_style
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(para_style.spacing_after)
            paragraph_format.line_spacing = para_style.line_height
        
        # Process paragraph content (handle inline elements)
        content = node.content
        
        # Handle links, images, emphasis in content
        self._add_formatted_text(paragraph, content, node.metadata)
    
    def _add_formatted_text(self, paragraph, content: str, metadata: Dict[str, Any]):
        """Add formatted text to paragraph, handling inline elements."""
        # Get base font settings
        base_font_size = None
        base_font_name = None
        
        if self.current_style_preset:
            para_style = self.current_style_preset.paragraph_style
            base_font_size = Pt(para_style.font.size)
            base_font_name = para_style.font.family
        
        # Use inline formatter to parse and apply formatting
        self.inline_formatter.format_text(
            paragraph, 
            content, 
            base_font_size=base_font_size,
            base_font_name=base_font_name
        )
    
    def _add_list(self, node: ASTNode, options: Dict[str, Any]):
        """Add list to document."""
        for list_item in node.children:
            if list_item.node_type == NodeType.LIST_ITEM:
                self._add_list_item(list_item, options)
    
    def _add_list_item(self, node: ASTNode, options: Dict[str, Any]):
        """Add list item to document."""
        is_ordered = node.metadata.get('ordered', False)
        content = node.content
        
        if is_ordered:
            paragraph = self.document.add_paragraph(content, style='List Number')
        else:
            paragraph = self.document.add_paragraph(content, style='List Bullet')
        
        # Apply list formatting
        paragraph_format = paragraph.paragraph_format
        level = node.level or 0
        paragraph_format.left_indent = Inches(0.25 * (level + 1))
    
    def _add_table(self, node: ASTNode, options: Dict[str, Any]):
        """Add table to document."""
        if not node.metadata:
            return
        
        headers = node.metadata.get('headers', [])
        rows = node.metadata.get('rows', [])
        
        if not headers:
            return
        
        # Create table
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = cell_data
    
    def _add_code_block(self, node: ASTNode, options: Dict[str, Any]):
        """Add code block to document."""
        content = node.content
        language = node.metadata.get('language')
        
        paragraph = self.document.add_paragraph()
        paragraph.style = 'No Spacing'
        
        # Apply code block styling
        if self.current_style_preset:
            code_style = self.current_style_preset.code_block_style
            
            # Set background color (requires custom XML)
            self._set_paragraph_background(paragraph, code_style.background_color)
            
            # Set border
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.25)
            paragraph_format.right_indent = Inches(0.25)
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
        
        # Add code content with monospace font
        run = paragraph.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        # Preserve line breaks
        content_lines = content.split('\n')
        for i, line in enumerate(content_lines):
            if i > 0:
                run.add_break()
            run.add_text(line)
    
    def _add_blockquote(self, node: ASTNode, options: Dict[str, Any]):
        """Add blockquote to document."""
        content = node.content
        
        paragraph = self.document.add_paragraph(content)
        paragraph.style = 'Quote'
        
        # Apply blockquote styling
        if self.current_style_preset:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.right_indent = Inches(0.5)
    
    def _add_horizontal_rule(self):
        """Add horizontal rule to document."""
        paragraph = self.document.add_paragraph()
        # Add a simple line using paragraph border
        paragraph_format = paragraph.paragraph_format
        # Use a simple approach - add a line of dashes
        run = paragraph.add_run('─' * 50)
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    def _add_image(self, node: ASTNode, options: Dict[str, Any]):
        """Add image to document."""
        image_src = node.attributes.get('src')
        if not image_src:
            return
        
        # Resolve image path
        image_path = Path(image_src)
        if not image_path.is_absolute():
            # Try relative to document directory or current working directory
            base_path = options.get('base_path', '.')
            image_path = Path(base_path) / image_path
        
        if not image_path.exists():
            if options.get('skip_missing_images', False):
                return
            raise FileError(
                f"Image file not found: {image_src}",
                file_path=str(image_path),
                operation="read"
            )
        
        try:
            # Add image to document
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            
            # Set image size (default: maintain aspect ratio, max width 6 inches)
            width = options.get('image_width', Inches(6))
            run.add_picture(str(image_path), width=width)
        
        except Exception as e:
            raise FileError(
                f"Failed to add image: {str(e)}",
                file_path=str(image_path),
                operation="read"
            ) from e
    
    def _add_bookmark(self, paragraph, name: str):
        """Add bookmark to paragraph for navigation."""
        # Clean bookmark name (Word bookmarks must be valid identifiers)
        bookmark_name = self._clean_bookmark_name(name)
        
        # Add bookmark using XML
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)
        run._element.getparent().insert(0, bookmark_start)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        run._element.getparent().append(bookmark_end)
    
    def _add_table_of_contents(self):
        """Add table of contents to document."""
        # Insert TOC at the beginning
        if not self.document.paragraphs:
            self.document.add_paragraph()
        
        # Add TOC heading
        toc_paragraph = self.document.paragraphs[0]
        toc_paragraph.insert_paragraph_before("Table of Contents")
        toc_paragraph = self.document.paragraphs[0]
        toc_paragraph.style = 'Heading 1'
        
        # Add TOC field using Word field codes
        # Create a paragraph for TOC
        toc_content_para = self.document.add_paragraph()
        
        # Insert TOC field (requires custom XML)
        # Full implementation would use Word's TOC field code
        # For now, we'll create a placeholder that Word can update
        run = toc_content_para.add_run()
        
        # Create TOC field element
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # Add page break after TOC
        self.document.add_page_break()
    
    def _set_paragraph_background(self, paragraph, color: str):
        """Set paragraph background color (requires custom XML)."""
        # This requires custom XML manipulation
        # Simplified implementation - full version would set shading
        pass
    
    def _parse_color(self, color_str: str) -> Any:
        """Parse color string to RGBColor."""
        if not DOCX_AVAILABLE:
            return None
        
        if color_str.startswith('#'):
            color_str = color_str[1:]
        
        if len(color_str) == 6:
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            return RGBColor(r, g, b)
        
        return RGBColor(0, 0, 0)
    
    def _clean_bookmark_name(self, name: str) -> str:
        """Clean bookmark name to be valid Word bookmark identifier."""
        # Remove invalid characters
        import re
        cleaned = re.sub(r'[^a-zA-Z0-9_]', '_', name)
        # Ensure it starts with a letter
        if cleaned and not cleaned[0].isalpha():
            cleaned = 'B' + cleaned
        return cleaned[:40]  # Limit length
    
    def _find_first_heading(self, node: ASTNode, level: int) -> Optional[ASTNode]:
        """Find first heading of specified level."""
        if node.node_type == NodeType.HEADING and node.level == level:
            return node
        
        for child in node.children:
            result = self._find_first_heading(child, level)
            if result:
                return result
        
        return None

